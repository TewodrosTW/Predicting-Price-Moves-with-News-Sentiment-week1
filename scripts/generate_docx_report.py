import os
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import nltk
from nltk.sentiment import SentimentIntensityAnalyzer
from scipy.stats import spearmanr

# For DOCX
from docx import Document
from docx.shared import Inches

# Paths
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
news_path = os.path.join(ROOT, 'data', 'raw_analyst_ratings.csv')
stock_path = os.path.join(ROOT, 'stock_data', 'AAPL.csv')
reports_dir = os.path.join(ROOT, 'reports')
os.makedirs(reports_dir, exist_ok=True)

if not os.path.exists(news_path):
    raise FileNotFoundError(f"News file not found: {news_path}")
if not os.path.exists(stock_path):
    raise FileNotFoundError(f"Stock file not found: {stock_path}")

# Ensure VADER
try:
    nltk.data.find('sentiment/vader_lexicon.zip')
except Exception:
    nltk.download('vader_lexicon')

sia = SentimentIntensityAnalyzer()

# Load & process news
news_df = pd.read_csv(news_path)
text_cols = [c for c in news_df.columns if c.lower() in ('headline','title','text','summary','article','news')]
if not text_cols:
    candidates = [c for c in news_df.columns if news_df[c].dtype == 'object']
    if candidates:
        text_cols = [max(candidates, key=lambda c: news_df[c].astype(str).map(len).mean())]
    else:
        raise KeyError('No text column detected in news CSV')
text_col = text_cols[0]

date_cols = [c for c in news_df.columns if c.lower() in ('date','publish_date','datetime','created_at','timestamp')]
if not date_cols:
    raise KeyError('No date-like column found in news CSV')
date_col = date_cols[0]
news_df[date_col] = pd.to_datetime(news_df[date_col], errors='coerce')
news_df = news_df.dropna(subset=[date_col, text_col])
news_df['publish_date'] = news_df[date_col].dt.date
news_df['sentiment'] = news_df[text_col].astype(str).map(lambda s: sia.polarity_scores(s)['compound'])

daily_sentiment = news_df.groupby('publish_date')['sentiment'].mean().rename('sentiment_mean')

# Load & process stock
sd = pd.read_csv(stock_path)
if 'Date' in sd.columns:
    sd['Date'] = pd.to_datetime(sd['Date'], errors='coerce')
elif 'date' in sd.columns:
    sd['date'] = pd.to_datetime(sd['date'], errors='coerce')
    sd = sd.rename(columns={'date':'Date'})
else:
    raise KeyError('No Date column in stock CSV')

sd = sd.set_index('Date').sort_index()
if 'Close' not in sd.columns:
    for alt in ['Adj Close','Adj_Close','AdjClose','close']:
        if alt in sd.columns:
            sd = sd.rename(columns={alt:'Close'})
            break
sd['Daily_Return'] = sd['Close'].pct_change()
stock_daily = sd['Daily_Return'].resample('D').mean().to_frame('stock_return')
stock_daily['date'] = stock_daily.index.date
stock_daily = stock_daily.set_index('date')

# Merge
merged = daily_sentiment.to_frame().join(stock_daily, how='inner').dropna()
merged = merged.sort_index()

# Stats
pearson = merged['sentiment_mean'].corr(merged['stock_return'])
spear, pval = spearmanr(merged['sentiment_mean'], merged['stock_return'])
lag_pearson = merged['sentiment_mean'].corr(merged['stock_return'].shift(-1))

# Rolling correlation
rolling_window = 90
rolling_corr = merged['sentiment_mean'].rolling(window=rolling_window).corr(merged['stock_return'])

# Generate figures
scatter_png = os.path.join(reports_dir, 'scatter.png')
fig, ax = plt.subplots(figsize=(8,6))
ax.scatter(merged['sentiment_mean'], merged['stock_return'], alpha=0.4)
ax.set_xlabel('Daily Mean Sentiment (compound)')
ax.set_ylabel('Daily Stock Return')
ax.set_title('Sentiment vs Daily Return')
ax.grid(True)
fig.tight_layout()
fig.savefig(scatter_png)
plt.close(fig)

timeseries_png = os.path.join(reports_dir, 'timeseries.png')
fig, (ax1, ax2) = plt.subplots(2,1, figsize=(10,8), sharex=True)
ax1.plot(merged.index, merged['sentiment_mean'], color='C0')
ax1.set_ylabel('Sentiment')
ax1.set_title('Daily Mean Sentiment')
ax2.plot(merged.index, merged['stock_return'], color='C1')
ax2.set_ylabel('Daily Return')
ax2.set_title('Daily Stock Return')
fig.tight_layout()
fig.savefig(timeseries_png)
plt.close(fig)

rolling_png = os.path.join(reports_dir, 'rolling_corr.png')
fig, ax = plt.subplots(figsize=(10,4))
ax.plot(merged.index, rolling_corr, label=f'Rolling {rolling_window}-day correlation')
ax.axhline(0, color='k', linewidth=0.5)
ax.set_ylabel('Rolling correlation')
ax.set_title(f'Rolling correlation ({rolling_window} days)')
ax.grid(True)
fig.tight_layout()
fig.savefig(rolling_png)
plt.close(fig)

# Create DOCX
docx_path = os.path.join(reports_dir, 'final_report.docx')
if os.path.exists(docx_path):
    os.remove(docx_path)

doc = Document()

doc.add_heading('Predicting Price Moves with News Sentiment — Final Report', level=1)

doc.add_heading('Executive Summary', level=2)
doc.add_paragraph(
    'This project investigates the relationship between news sentiment and stock price movements. '
    'We processed news headlines, computed per-article sentiment using NLTK VADER (compound score), aggregated daily sentiment, computed daily stock returns, aligned both datasets by date, and measured correlations. '
)

# Implementation overview

doc.add_heading('Implementation & Methodology', level=2)
doc.add_paragraph('Tools and Libraries: Python, pandas, matplotlib, NLTK (VADER), scipy, python-docx.')

doc.add_heading('Data Preparation', level=3)
doc.add_paragraph(
    'News data was loaded from: ' + news_path + '. The code detects a suitable text column (headline/title/text) and a date column, normalizes timestamps to dates, and drops invalid rows. '
    'Stock data was loaded from: ' + stock_path + '. Date and Close columns were normalized and daily returns computed as pct_change().' )

# Key steps

doc.add_heading('Key Steps Taken', level=3)
doc.add_paragraph('- Date alignment: normalized timestamps and aggregated news to calendar days.')
doc.add_paragraph('- Sentiment analysis: per-article compound score using VADER.')
doc.add_paragraph('- Aggregation: daily mean sentiment used as primary signal.')
doc.add_paragraph('- Stock movement: daily percent returns from Close price.')
doc.add_paragraph('- Correlation analysis: Pearson and Spearman correlations, and rolling correlation window (90 days).')

# Findings

doc.add_heading('Key Findings', level=2)
doc.add_paragraph(f'- Days with both news and returns: {len(merged)}')
doc.add_paragraph(f'- Pearson correlation (daily mean sentiment vs daily return): {pearson:.6f}')
doc.add_paragraph(f'- Spearman correlation: {spear:.6f} (p-value: {pval:.6g})')
doc.add_paragraph(f'- Lagged Pearson (sentiment -> next day return): {lag_pearson:.6f}')

# Insert figures

doc.add_heading('Figures', level=2)
doc.add_paragraph('Scatter: daily mean sentiment vs daily stock return')
doc.add_picture(scatter_png, width=Inches(6))

doc.add_paragraph('Time series: daily mean sentiment (top) and daily returns (bottom)')
doc.add_picture(timeseries_png, width=Inches(6))

doc.add_paragraph(f'Rolling {rolling_window}-day correlation between sentiment and returns')
doc.add_picture(rolling_png, width=Inches(6))

# Challenges & Lessons

doc.add_heading('Challenges and Lessons Learned', level=2)
doc.add_paragraph('- Library compatibility: Some libraries like pandas_ta require numba and specific numpy versions; to keep the environment stable we used pandas-only computations for indicators.')
doc.add_paragraph('- Date alignment: mapping article timestamps to trading days requires domain rules (e.g., post-market news may map to next trading day).')
doc.add_paragraph('- Sentiment limitations: VADER is tuned for social media; headline sentiment may require more sophisticated models for nuance.')

# Conclusions & Future Work

doc.add_heading('Conclusions and Future Work', level=2)
doc.add_paragraph(
    'The correlation between daily mean news sentiment and daily stock returns is small (near zero in our results). '
    'This suggests that simple daily-mean headline sentiment has limited predictive power for daily returns in this dataset. '
)

doc.add_paragraph('Future improvements: use weighted sentiment, intraday alignment, transformer-based sentiment models, event detection, and more robust statistical testing (p-values, significance).')

# Appendix

doc.add_heading('Appendix', level=2)
doc.add_paragraph('Files created:')
doc.add_paragraph('- notebooks/task2_stock_eda.ipynb (pandas-only EDA)')
doc.add_paragraph('- notebooks/task3_correlation.ipynb (correlation analysis notebook)')
doc.add_paragraph('- reports/final_report.docx (this document)')
doc.add_paragraph('- reports/task3_report.pdf (earlier PDF report)')

doc.add_heading('How to reproduce', level=3)
doc.add_paragraph('1. Activate the virtual environment: `.\.venv\Scripts\Activate.ps1`')
doc.add_paragraph('2. Install dependencies: `python -m pip install -r requirements.txt` (or `python -m pip install pandas matplotlib nltk scipy python-docx`)')
doc.add_paragraph('3. Open and run the notebooks in order: task2_stock_eda.ipynb then task3_correlation.ipynb')

# Save

doc.save(docx_path)
print('DOCX report saved to:', docx_path)
